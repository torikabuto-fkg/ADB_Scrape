
import os
import glob
import img2pdf
import paddle
from paddleocr import PaddleOCR
from natsort import natsorted
from docx import Document
from tqdm import tqdm
import logging

# PaddleOCRの冗長なログを抑制
logging.getLogger("ppocr").setLevel(logging.WARNING)

# --- 設定エリア ---
# 画像が入っているフォルダのパス（末尾の / はあってもなくても動きます）
IMAGE_DIR = "./hokuto_scrapes/"
# 結果を出力するフォルダ
OUTPUT_DIR = "./"

# 出力ファイル名
PDF_FILENAME = "_reviews.pdf"
DOCX_FILENAME = "_reviews.docx"
TXT_FILENAME = "_raw_text.txt"

# OCRの信頼度閾値（ノイズ除去用: 0.0〜1.0）
CONFIDENCE_THRESHOLD = 0.6 

def main():
    # 出力ディレクトリがなければ作成
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    # 1. CPU設定
    # エラー回避のため、明示的にCPUを指定します
    paddle.set_device('cpu') 
    print("✅ CPUモードで動作します (安定版)")

    # 2. 画像リスト取得
    # natsortedを使うことで page_1.png -> page_2.png -> page_10.png の順に正しく並びます
    img_paths = natsorted(glob.glob(os.path.join(IMAGE_DIR, "*.png")))
    
    if not img_paths:
        print(f"エラー: '{IMAGE_DIR}' に画像が見つかりません。パスを確認してください。")
        return

    print(f"対象画像: {len(img_paths)}枚")

    # =========================================================
    # Phase 1: PDF作成 (自炊・閲覧用)
    # =========================================================
    print("\n[1/3] PDFを作成しています...")
    pdf_path = os.path.join(OUTPUT_DIR, PDF_FILENAME)
    
    try:
        with open(pdf_path, "wb") as f:
            # 画像をそのままPDFに変換（劣化なし）
            f.write(img2pdf.convert(img_paths))
        print(f"  -> PDF保存完了: {pdf_path}")
    except Exception as e:
        print(f"  -> PDF作成エラー: {e}")

    # =========================================================
    # Phase 2: PaddleOCR実行 (テキスト化)
    # =========================================================
    print("\n[2/3] PaddleOCRで解析しています...")
    
    # PaddleOCRの初期化 (CPU設定)
    # use_angle_cls=False: 向き補正なし（スクショは縦向き固定のため不要＆高速化）
    # use_gpu=False: CPU強制
    ocr = PaddleOCR(
        use_angle_cls=False,
        lang='japan',
        use_gpu=False
    )
    
    full_text_data = []

    # プログレスバーを表示しながら処理
    for img_path in tqdm(img_paths, desc="OCR Progress"):
        file_name = os.path.basename(img_path)
        
        # OCR実行 (cls=TrueでもinitでOFFにしているのでスキップされます)
        result = ocr.ocr(img_path, cls=True)
        
        page_text = []
        
        # 結果がある場合のみ処理
        if result and result[0]:
            for line in result[0]:
                text = line[1][0]
                score = line[1][1]
                
                # 信頼度が低い文字（ノイズ）を除外
                if score >= CONFIDENCE_THRESHOLD:
                    page_text.append(text)
        
        # ページごとのデータをリストに格納
        full_text_data.append({
            "filename": file_name,
            "content": "\n".join(page_text)
        })

    # =========================================================
    # Phase 3: ファイル出力 (Word & Text)
    # =========================================================
    print("\n[3/3] 結果をファイルに書き出しています...")

    # A. テキストファイル出力 (AI分析・LLM入力用)
    txt_path = os.path.join(OUTPUT_DIR, TXT_FILENAME)
    with open(txt_path, "w", encoding="utf-8") as f:
        for page in full_text_data:
            f.write(f"--- Page: {page['filename']} ---\n")
            f.write(page['content'])
            f.write("\n\n")
    
    # B. Wordファイル出力 (閲覧・編集用)
    docx_path = os.path.join(OUTPUT_DIR, DOCX_FILENAME)
    doc = Document()
    doc.add_heading('Hokuto Reviews (PaddleOCR Result)', 0)
    
    for page in full_text_data:
        doc.add_heading(f"Page: {page['filename']}", level=2)
        if page['content'].strip():
            doc.add_paragraph(page['content'])
        else:
            doc.add_paragraph("(テキストなし)")
            
    doc.save(docx_path)
    
    print(f"\n✅ 完了: {OUTPUT_DIR} を確認してください")

if __name__ == "__main__":
    main()
